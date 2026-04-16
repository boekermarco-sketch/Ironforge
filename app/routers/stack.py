import io
from fastapi import APIRouter, Request, Depends, Form, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.responses import RedirectResponse, StreamingResponse
from sqlalchemy.orm import Session
from datetime import date
from pathlib import Path

from app.database import get_db
from app.models import Stack, Substance, DoseEvent, StackChangeLog

router = APIRouter()
BASE_DIR = Path(__file__).resolve().parent.parent.parent
templates = Jinja2Templates(directory=BASE_DIR / "templates")


def _log(db: Session, object_type: str, object_id: int, action: str,
         summary: str, field_name: str = None, old_value=None,
         new_value=None, reason: str = None):
    """Schreibt einen Eintrag in stack_changelog."""
    db.add(StackChangeLog(
        object_type=object_type,
        object_id=object_id,
        action=action,
        field_name=field_name,
        old_value=str(old_value) if old_value is not None else None,
        new_value=str(new_value) if new_value is not None else None,
        summary=summary,
        reason=reason or None,
    ))


@router.get("/", include_in_schema=False)
async def stack_overview(request: Request, db: Session = Depends(get_db)):
    today = date.today()
    stacks = db.query(Stack).order_by(Stack.start_date.desc()).all()

    active_stack = next((s for s in stacks if s.status == "aktiv"), None)
    active_doses = []
    if active_stack:
        rows = (
            db.query(DoseEvent, Substance)
            .join(Substance, DoseEvent.substance_id == Substance.id)
            .filter(DoseEvent.stack_id == active_stack.id)
            .filter(DoseEvent.start_date <= today)
            .filter((DoseEvent.end_date == None) | (DoseEvent.end_date >= today))
            .order_by(Substance.category, Substance.name)
            .all()
        )
        # Gruppiert nach Kategorie
        grouped: dict[str, list] = {}
        for de, sub in rows:
            grouped.setdefault(sub.category, []).append({"dose": de, "substance": sub})
        active_doses = grouped

    substances = db.query(Substance).filter(Substance.active == True).order_by(Substance.category, Substance.name).all()
    recent_changes = (
        db.query(StackChangeLog)
        .order_by(StackChangeLog.changed_at.desc())
        .limit(30)
        .all()
    )

    # Zusätzlich nach Zeitslot gruppieren für die neue Ansicht
    SLOTS = [
        "07:00 – Nüchtern",
        "10–11 Uhr",
        "17–18 Uhr",
        "21–22 Uhr",
        "Injektionen",
        "Sonstige",
    ]

    def _slot(timing: str) -> str:
        t = (timing or "").lower()
        if any(k in t for k in ("subkutan", "intramuskulär", "intramuskulaer", "mi abend", "so morgen", "donnerstag")):
            return "Injektionen"
        if "07" in t or "nüchtern" in t or ("morgen" in t and "so morgen" not in t and "donnerstag" not in t):
            return "07:00 – Nüchtern"
        if "10" in t or "11" in t:
            return "10–11 Uhr"
        if "17" in t or "18" in t:
            return "17–18 Uhr"
        if "21" in t or "22" in t or "nacht" in t:
            return "21–22 Uhr"
        return "Sonstige"

    doses_by_slot: dict[str, list] = {s: [] for s in SLOTS}
    if active_stack:
        rows_all = (
            db.query(DoseEvent, Substance)
            .join(Substance, DoseEvent.substance_id == Substance.id)
            .filter(DoseEvent.stack_id == active_stack.id)
            .filter(DoseEvent.start_date <= today)
            .filter((DoseEvent.end_date == None) | (DoseEvent.end_date >= today))
            .order_by(Substance.category, Substance.name)
            .all()
        )
        for de, sub in rows_all:
            slot = _slot(de.timing or "")
            doses_by_slot[slot].append({"dose": de, "substance": sub})

    return templates.TemplateResponse("stack.html", {
        "request": request,
        "stacks": stacks,
        "active_stack": active_stack,
        "active_doses": active_doses,
        "doses_by_slot": doses_by_slot,
        "slot_order": SLOTS,
        "substances": substances,
        "today": today,
        "recent_changes": recent_changes,
    })


@router.post("/{stack_id}/bearbeiten", include_in_schema=False)
async def edit_stack(
    stack_id: int,
    name: str = Form(...),
    goal: str = Form(""),
    start_date: str = Form(...),
    end_date: str = Form(""),
    status: str = Form("aktiv"),
    notes: str = Form(""),
    db: Session = Depends(get_db)
):
    stack = db.query(Stack).filter(Stack.id == stack_id).first()
    if not stack:
        raise HTTPException(status_code=404, detail="Stack nicht gefunden")

    changed_fields = []
    for field, old, new in [
        ("name", stack.name, name),
        ("goal", stack.goal, goal or None),
        ("status", stack.status, status),
        ("start_date", str(stack.start_date), start_date),
        ("end_date", str(stack.end_date) if stack.end_date else None, end_date or None),
    ]:
        if str(old or "") != str(new or ""):
            changed_fields.append(field)
            _log(db, "stack", stack_id, "updated",
                 f"Stack '{name}': {field} geändert",
                 field_name=field, old_value=old, new_value=new)

    stack.name = name
    stack.goal = goal or None
    stack.start_date = date.fromisoformat(start_date)
    stack.end_date = date.fromisoformat(end_date) if end_date else None
    stack.status = status
    stack.notes = notes or None
    if not changed_fields:
        _log(db, "stack", stack_id, "updated", f"Stack '{name}' gespeichert (keine Änderung erkannt)")
    db.commit()
    return RedirectResponse("/stack/", status_code=303)


@router.post("/substanz/neu", include_in_schema=False)
async def add_substance(
    request: Request,
    name: str = Form(...),
    category: str = Form(...),
    route: str = Form("oral"),
    default_unit: str = Form("mg"),
    description: str = Form(""),
    db: Session = Depends(get_db)
):
    existing = db.query(Substance).filter(Substance.name == name).first()
    if existing:
        raise HTTPException(status_code=400, detail="Substanz existiert bereits")
    db.add(Substance(name=name, category=category, route=route, default_unit=default_unit, description=description))
    db.commit()
    return RedirectResponse("/stack/", status_code=303)


@router.post("/dosis/neu", include_in_schema=False)
async def add_dose_event(
    substance_id: int = Form(...),
    stack_id: int = Form(...),
    dose_amount: float = Form(...),
    dose_unit: str = Form("mg"),
    frequency: str = Form(""),
    timing: str = Form(""),
    start_date: str = Form(...),
    change_reason: str = Form(""),
    notes: str = Form(""),
    db: Session = Depends(get_db)
):
    new_de = DoseEvent(
        substance_id=substance_id,
        stack_id=stack_id,
        dose_amount=dose_amount,
        dose_unit=dose_unit,
        frequency=frequency,
        timing=timing,
        start_date=date.fromisoformat(start_date),
        change_reason=change_reason or None,
        notes=notes or None,
    )
    db.add(new_de)
    db.flush()  # ID zuweisen ohne zu committen

    substance = db.query(Substance).filter(Substance.id == substance_id).first()
    sub_name = substance.name if substance else f"#{substance_id}"
    _log(db, "dose_event", new_de.id, "created",
         f"{sub_name}: {dose_amount}{dose_unit} hinzugefügt",
         reason=change_reason or None)
    db.commit()
    return RedirectResponse("/stack/", status_code=303)


@router.post("/dosis/{dose_id}/bearbeiten", include_in_schema=False)
async def edit_dose_event(
    dose_id: int,
    dose_amount: float = Form(...),
    dose_unit: str = Form("mg"),
    frequency: str = Form(""),
    timing: str = Form(""),
    notes: str = Form(""),
    db: Session = Depends(get_db)
):
    de = db.query(DoseEvent).filter(DoseEvent.id == dose_id).first()
    if de:
        substance = db.query(Substance).filter(Substance.id == de.substance_id).first()
        sub_name = substance.name if substance else f"#{de.substance_id}"

        if de.dose_amount != dose_amount or de.dose_unit != dose_unit:
            _log(db, "dose_event", dose_id, "updated",
                 f"{sub_name}: {de.dose_amount}{de.dose_unit} → {dose_amount}{dose_unit}",
                 field_name="dose_amount",
                 old_value=f"{de.dose_amount}{de.dose_unit}",
                 new_value=f"{dose_amount}{dose_unit}")
        if de.frequency != frequency:
            _log(db, "dose_event", dose_id, "updated",
                 f"{sub_name}: Frequenz '{de.frequency}' → '{frequency}'",
                 field_name="frequency", old_value=de.frequency, new_value=frequency)
        if de.timing != timing:
            _log(db, "dose_event", dose_id, "updated",
                 f"{sub_name}: Timing geändert",
                 field_name="timing", old_value=de.timing, new_value=timing)

        de.dose_amount = dose_amount
        de.dose_unit = dose_unit
        de.frequency = frequency
        de.timing = timing
        de.notes = notes or None
        db.commit()
    return RedirectResponse("/stack/", status_code=303)


@router.post("/dosis/{dose_id}/abschliessen", include_in_schema=False)
async def end_dose_event(
    dose_id: int,
    end_date: str = Form(default=None),
    db: Session = Depends(get_db)
):
    de = db.query(DoseEvent).filter(DoseEvent.id == dose_id).first()
    if de:
        closed = date.fromisoformat(end_date) if end_date else date.today()
        substance = db.query(Substance).filter(Substance.id == de.substance_id).first()
        sub_name = substance.name if substance else f"#{de.substance_id}"
        _log(db, "dose_event", dose_id, "updated",
             f"{sub_name}: abgeschlossen ab {closed.strftime('%d.%m.%Y')}",
             field_name="end_date", old_value=None, new_value=str(closed))
        de.end_date = closed
        db.commit()
    return RedirectResponse("/stack/", status_code=303)


# ── Stack-Plan als DOCX herunterladen ─────────────────────────────────────────

_SLOTS = [
    ("07:00 – Nüchtern",          lambda t: "07" in t or "nüchtern" in t.lower() or "morgens" in t.lower()),
    ("10–11 Uhr – Mit Frühstück", lambda t: "10" in t or "11" in t),
    ("17–18 Uhr – Mit Mahlzeit",  lambda t: "17" in t or "18" in t),
    ("21–22 Uhr – Zur Nacht",     lambda t: "21" in t or "22" in t or "nacht" in t.lower()),
    ("Injektionen (Mi + So)",     lambda t: any(k in t.lower() for k in ("subkutan", "intramuskulär", "intramuskulaer", "mi abend", "so morgen", "donnerstag"))),
]

_CATEGORY_ORDER = ["Steroid", "Hormon", "Peptid", "Medikament", "Supplement"]


def _slot_for(timing: str) -> str:
    t = timing or ""
    # Injektionen zuerst prüfen
    if any(k in t.lower() for k in ("subkutan", "intramuskulär", "intramuskulaer", "mi abend", "so morgen", "donnerstag")):
        return "Injektionen (Mi + So)"
    for label, check in _SLOTS[:-1]:
        if check(t):
            return label
    return "Sonstige"


def _generate_stack_docx(db: Session) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    today = date.today()
    stack = db.query(Stack).filter(Stack.status == "aktiv").first()
    if not stack:
        stack = db.query(Stack).order_by(Stack.start_date.desc()).first()

    rows = (
        db.query(DoseEvent, Substance)
        .join(Substance, DoseEvent.substance_id == Substance.id)
        .filter(DoseEvent.stack_id == stack.id)
        .filter(DoseEvent.start_date <= today)
        .filter((DoseEvent.end_date == None) | (DoseEvent.end_date >= today))
        .order_by(Substance.category, Substance.name)
        .all()
    )

    # Gruppieren nach Slot
    slots: dict[str, list] = {s[0]: [] for s in _SLOTS}
    slots["Sonstige"] = []
    for ev, sub in rows:
        slot = _slot_for(ev.timing or "")
        slots.setdefault(slot, []).append((ev, sub))

    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titel
    title = doc.add_heading("Supplement & Medikamenten-Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_title = doc.add_paragraph(f"Marco Böker · Stand: {today.strftime('%d.%m.%Y')} · {stack.name}")
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title.runs[0].font.size = Pt(10)
    sub_title.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Blast-Info
    blast_start = date(2026, 2, 23)
    if today >= blast_start:
        blast_day = (today - blast_start).days + 1
        blast_week = ((today - blast_start).days // 7) + 1
        blast_p = doc.add_paragraph(f"Blast-Tag {blast_day} · Woche {blast_week} von 16")
        blast_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        blast_p.runs[0].font.size = Pt(9)
        blast_p.runs[0].font.color.rgb = RGBColor(0x33, 0x99, 0x44)

    doc.add_paragraph()

    # Slots
    for slot_label, _ in _SLOTS:
        items = slots.get(slot_label, [])
        if not items:
            continue

        h = doc.add_heading(slot_label, level=1)
        h.runs[0].font.size = Pt(12)

        # Tabelle
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, txt in enumerate(["Substanz", "Dosis", "Frequenz", "Hinweis/Notiz"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

        # Sortierung innerhalb Slot nach Kategorie
        items_sorted = sorted(items, key=lambda x: (_CATEGORY_ORDER.index(x[1].category) if x[1].category in _CATEGORY_ORDER else 99, x[1].name))

        for ev, sub in items_sorted:
            row = tbl.add_row().cells
            row[0].text = sub.name
            row[0].paragraphs[0].runs[0].font.size = Pt(9)
            dose_str = f"{ev.dose_amount:g} {ev.dose_unit}"
            row[1].text = dose_str
            row[1].paragraphs[0].runs[0].font.size = Pt(9)
            row[2].text = ev.frequency or ""
            row[2].paragraphs[0].runs[0].font.size = Pt(9)
            note = ev.notes or ev.timing or ""
            row[3].text = note
            row[3].paragraphs[0].runs[0].font.size = Pt(8)

        doc.add_paragraph()

    # Sonstige
    if slots.get("Sonstige"):
        h = doc.add_heading("Sonstige / kein fester Zeitslot", level=1)
        h.runs[0].font.size = Pt(12)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, txt in enumerate(["Substanz", "Dosis", "Frequenz", "Timing / Notiz"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
        for ev, sub in slots["Sonstige"]:
            row = tbl.add_row().cells
            row[0].text = sub.name
            row[1].text = f"{ev.dose_amount:g} {ev.dose_unit}"
            row[2].text = ev.frequency or ""
            row[3].text = ev.timing or ev.notes or ""

    # Hinweis E2
    doc.add_paragraph()
    note_p = doc.add_paragraph(
        "E2-Hinweis: Exemestan läuft seit 02.04.2026. "
        "10 Tage täglich 25mg (bis 11.04.), danach alternierend 12,5mg / 25mg je Tag. "
        "Ziel: E2-Senkung von 99 ng/L → Normbereich (20–40 ng/L)."
    )
    note_p.runs[0].font.size = Pt(8)
    note_p.runs[0].font.color.rgb = RGBColor(0xAA, 0x55, 0x00)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/download", include_in_schema=False)
async def download_stack_plan(db: Session = Depends(get_db)):
    """Generiert den aktuellen Stack-Plan als DOCX und liefert ihn zum Download."""
    docx_bytes = _generate_stack_docx(db)
    filename = f"Stack_Plan_{date.today().strftime('%Y-%m-%d')}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
